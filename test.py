from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    # 添加标题
    heading = doc.add_heading(text, level=level)
    # 居中对齐
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 创建新的Word文档
doc = Document()

# 添加不同级别的标题
add_heading(doc, 'Main Title', level=0)  # Heading 1
add_heading(doc, 'Sub Title 1', level=1)  # Heading 2
add_heading(doc, 'Sub Title 2', level=2)  # Heading 3

# 保存文档
doc.save('document_with_headings.docx')

